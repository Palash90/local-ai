#!/usr/bin/env python3
import http.server, json, os, uuid, base64, mimetypes, requests, subprocess, time, random, threading, sys, io, tempfile, queue as _queue_mod, traceback

sys.stdout.reconfigure(line_buffering=True)  # noqa
from datetime import datetime
from urllib.parse import urlparse, parse_qs
from concurrent.futures import ThreadPoolExecutor

LLAMA_BASE = "http://localhost:8081"
LLAMA_URL = f"{LLAMA_BASE}/v1/chat/completions"

VENV_PYTHON = os.path.expanduser("~/local-ai/ComfyUI/venv/bin/python")
COMFYUI_DIR = os.path.expanduser("~/local-ai/ComfyUI")
SEARXNG_URL = os.environ.get("SEARXNG_URL", "http://localhost:8080/search")
COMFYUI_URL = "http://localhost:8188"
HOST = os.environ.get("CHAT_HOST", "0.0.0.0")
PORT = 3001
REASONING_BUDGET = 4096

with open(os.path.expanduser("~/local-ai-files/model.txt"), "r") as file:
    MODEL_ID = file.read()

import sys

sys.path.insert(0, COMFYUI_DIR)
COMFYUI_OUTPUT = os.path.expanduser("~/local-ai-files/ComfyUI/output")
UPLOADS_DIR = os.path.expanduser("~/local-ai-files/uploads")
LLAMA_SERVER_PATH = os.path.expanduser("~/local-ai/llama.cpp/build/bin/llama-server")
LLAMA_QWEN_NGL = "12"
LLAMA_GEMMA_NGL = "99"
LLAMA_SERVER_ARGS = [
    "--host", "0.0.0.0",
    "--port", "8081",
    "--models-dir", os.path.expanduser("~/local-ai-files/my-models/"),
    "--jinja",
    
    # GPU / VRAM Allocations
    "--n-gpu-layers", "99",
    "-fa", "on",  # Flash attention lowers VRAM footprint
    "--ctx-size", "32768",  # 32k context; KV cache quantized to q8_0 to fit VRAM
    #"--no-kv-offload",
    "-ctk", "q8_0",            # Quantize Key cache to 8-bit (saves 50% VRAM)
    "-ctv", "q8_0",            # Quantize Value cache to 8-bit (saves 50% VRAM)
    
    # Reasoning & Thinking Limits
    "--reasoning-budget", str(REASONING_BUDGET),
    "--reasoning-budget-message", "Reasoning limit reached, summarize final answer.",
    
    # Gemma 4 Sampling Preset
    "--temp", "1.0",
    "--top-p", "0.95",
    "--top-k", "64",
    "--min-p", "0.0",
    "--repeat-penalty", "1.0"
]

SESSIONS_FILE = os.path.expanduser("~/local-ai-files/sessions.json")
IMG_PATH = os.path.expanduser("~/local-ai-files/ComfyUI/output")
COMFYUI_INPUT = os.path.expanduser("~/local-ai-files/ComfyUI/input")
PROMPT_PATH = os.path.expanduser("~/local-ai-files/sys_prompt.txt")
USERS_FILE = os.path.expanduser("~/local-ai-files/users.json")
TASKS_DB = os.path.expanduser("~/local-ai-files/tasks.db")
IMAGE_TOKEN_COST = 1200
AUDIO_TOKEN_COST = 800
PER_MESSAGE_OVERHEAD = 4

with open(
    os.path.expanduser("~/local-ai-files/models.json"), "r", encoding="utf-8"
) as file:
    IMAGE_MODELS = json.load(file)

TOOLS = [
        {
            "type": "function",
            "function": {
                "name": "web_search",
                "description": "Search the web for real-time/current information. Use this for weather, news, sports, stock prices, recent events, or any query where up-to-date data matters. Do NOT answer time-sensitive questions from memory — always search. The results contain snippets only; if the snippets are insufficient to answer the question fully, follow up with fetch_page to read the full content of the relevant page.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "The search query"},
                        "current_time": {"type": "string", "description": "Current date and time. Pass ONLY for time-sensitive queries (news, events, hours, etc.) where recency matters. Omit for direct-link lookups or general information."},
                        "current_location": {"type": "string", "description": "User's location. Pass ONLY for location-specific results (weather, local news, nearby places, events). If you don't know the user's location, call get_user_location first to obtain it. Do NOT guess or fabricate location."}
                    },
                    "required": ["query"],
                },
            },
        },
    {
        "type": "function",
        "function": {
            "name": "fetch_page",
            "description": "Fetch and read the full text content of a web page. Use this AFTER web_search when the search snippets are not enough to answer the question (e.g. you need details, data, or an article's body). Pass the full URL of the page to read.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "The full URL of the web page to fetch (must start with http:// or https://)."
                    }
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_image",
            "description": "Generate or draw an image. You MUST choose a style model.",
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {
                        "type": "string",
                        "description": "Detailed visual description of what to draw/generate.",
                    },
                    "negative_prompt": {
                        "type": "string",
                        "description": "Things to avoid in the image",
                    },
                    "model": {
                        "type": "string",
                        "enum": list(IMAGE_MODELS.keys()),
                        "description": "Art style to use. Options: "
                        + ", ".join(
                            [
                                f"'{k}' ({v['description']})"
                                for k, v in IMAGE_MODELS.items()
                            ]
                        ),
                    },
                },
                "required": ["prompt", "model"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "edit_image",
            "description": "Generic Img2Img image editor to modify, restyle, recolor, add elements, or transform existing or uploaded images.",
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {
                        "type": "string",
                        "description": "Complete description of what the edited image should look like.",
                    },
                    "negative_prompt": {
                        "type": "string",
                        "description": "Elements to exclude from the visual generation.",
                    },
                    "denoise": {
                        "type": "number",
                        "description": "Denoising value (0.1 to 1.0). Use 0.25-0.4 for subtle color/lighting changes, 0.45-0.65 for structural edits and object additions, and 0.7-0.85 for massive re-imaginings.",
                    },
                },
                "required": ["prompt", "denoise"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_user_location",
            "description": "Request the user's current geographical location. Call this ONLY when you need location for a location-specific query (weather, local news, nearby places, etc.) and you don't already have the user's location. Returns the user's city/area or 'denied' if they refuse.",
            "parameters": {
                "type": "object",
                "properties": {},
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": "Read the text content of an uploaded file (PDF, DOC, DOCX, XLS, XLSX). Call this when the user has attached a file and you need to read its content to answer their question. The file URL is provided in the user message as [FILE: url]. Pass that url as the file_url parameter.",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_url": {
                        "type": "string",
                        "description": "The file URL from the user message (the /uploads/... path)."
                    }
                },
                "required": ["file_url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "update_user_context",
            "description": "Store information about the current user that persists across conversations. Saves preferences, personal details, important facts, or anything the user should not need to repeat. This APPENDS to existing context — only add NEW information, do not repeat what was already saved.",
            "parameters": {
                "type": "object",
                "properties": {
                    "content": {
                        "type": "string",
                        "description": "The new information to append to the user's context. Keep it concise and focused on what's new.",
                    }
                },
                "required": ["content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "manage_tasks",
            "description": "Manage to-do tasks. Can create, update, complete, delete, list, or get task details. Use this when the user wants to track tasks, set reminders, or manage their to-do list.",
            "parameters": {
                "type": "object",
                "properties": {
                    "operation": {
                        "type": "string",
                        "enum": ["create", "update", "complete", "delete", "list", "get"],
                        "description": "The operation to perform.",
                    },
                    "task_id": {
                        "type": "string",
                        "description": "Required for update/complete/delete/get. The task ID.",
                    },
                    "title": {
                        "type": "string",
                        "description": "Required for create. Task title.",
                    },
                    "description": {
                        "type": "string",
                        "description": "Task description or details.",
                    },
                    "priority": {
                        "type": "string",
                        "enum": ["low", "medium", "high"],
                        "description": "Task priority (default: medium).",
                    },
                    "status": {
                        "type": "string",
                        "enum": ["pending", "in_progress", "completed", "cancelled"],
                        "description": "For update: new status.",
                    },
                    "due_date": {
                        "type": "string",
                        "description": "Due date in ISO format (e.g. 2026-08-15T17:00:00).",
                    },
                    "reminder_at": {
                        "type": "string",
                        "description": "Reminder time in ISO format. The system will notify about this task at the given time.",
                    },
                    "session_id": {
                        "type": "string",
                        "description": "Session ID to link this task to a conversation.",
                    },
                },
                "required": ["operation"],
            },
        },
    },
]

TOOLS_TOKEN_COST = len(json.dumps(TOOLS)) // 4

import sqlite3
_tasks_db_lock = threading.Lock()

def _init_tasks_db():
    with _tasks_db_lock:
        conn = sqlite3.connect(TASKS_DB)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS tasks (
                id TEXT PRIMARY KEY,
                user_id TEXT NOT NULL,
                title TEXT NOT NULL,
                description TEXT DEFAULT '',
                status TEXT DEFAULT 'pending',
                priority TEXT DEFAULT 'medium',
                due_date TEXT,
                session_id TEXT,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                reminder_at TEXT,
                reminded INTEGER DEFAULT 0
            )
        """)
        conn.commit()
        conn.close()

def _db_run(query, params=()):
    with _tasks_db_lock:
        conn = sqlite3.connect(TASKS_DB)
        conn.row_factory = sqlite3.Row
        try:
            cur = conn.execute(query, params)
            conn.commit()
            return cur
        finally:
            conn.close()

def _db_fetch(query, params=()):
    with _tasks_db_lock:
        conn = sqlite3.connect(TASKS_DB)
        conn.row_factory = sqlite3.Row
        try:
            cur = conn.execute(query, params)
            rows = [dict(r) for r in cur.fetchall()]
            return rows
        finally:
            conn.close()

def _db_fetch_one(query, params=()):
    rows = _db_fetch(query, params)
    return rows[0] if rows else None

_init_tasks_db()




with open(PROMPT_PATH, "r") as file:
    SYS_CONTENT = file.read()
model_list = "; ".join(f"{k}: {v['description']}" for k, v in IMAGE_MODELS.items())
SYS_CONTENT = SYS_CONTENT.replace("%model_list%", model_list)
SYS_CONTENT = SYS_CONTENT.replace("%_image_keys%", str(list(IMAGE_MODELS.keys())))

print("Prompt:\n", "*" * 80, "\n", SYS_CONTENT, "\n", "*" * 80)

def task_create(user_id, title, description="", priority="medium", due_date=None, session_id=None, reminder_at=None):
    tid = str(uuid.uuid4())
    now = datetime.now().isoformat()
    _db_run(
        "INSERT INTO tasks (id, user_id, title, description, status, priority, due_date, session_id, created_at, updated_at, reminder_at) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
        (tid, user_id, title, description, "pending", priority, due_date, session_id, now, now, reminder_at),
    )
    return _db_fetch_one("SELECT * FROM tasks WHERE id=?", (tid,))

def task_update(tid, user_id, **kwargs):
    fields = {k: v for k, v in kwargs.items() if v is not None}
    if not fields:
        return None
    fields["updated_at"] = datetime.now().isoformat()
    set_clause = ", ".join(f"{k}=?" for k in fields)
    vals = list(fields.values()) + [tid, user_id]
    _db_run(f"UPDATE tasks SET {set_clause} WHERE id=? AND user_id=?", vals)
    return _db_fetch_one("SELECT * FROM tasks WHERE id=?", (tid,))

def task_complete(tid, user_id):
    now = datetime.now().isoformat()
    _db_run("UPDATE tasks SET status='completed', updated_at=? WHERE id=? AND user_id=?", (now, tid, user_id))
    return _db_fetch_one("SELECT * FROM tasks WHERE id=?", (tid,))

def task_delete(tid, user_id):
    _db_run("DELETE FROM tasks WHERE id=? AND user_id=?", (tid, user_id))

def task_list(user_id, status=None):
    if status:
        return _db_fetch("SELECT * FROM tasks WHERE user_id=? AND status=? ORDER BY due_date IS NULL, due_date ASC, created_at DESC", (user_id, status))
    return _db_fetch("SELECT * FROM tasks WHERE user_id=? ORDER BY due_date IS NULL, due_date ASC, created_at DESC", (user_id,))

def task_get(tid, user_id):
    return _db_fetch_one("SELECT * FROM tasks WHERE id=? AND user_id=?", (tid, user_id))

def handle_task_tool(user_id, args):
    op = args.get("operation", "")
    if op == "create":
        t = task_create(user_id, args["title"], args.get("description", ""), args.get("priority", "medium"), args.get("due_date"), args.get("session_id"), args.get("reminder_at"))
        return json.dumps({"ok": True, "task": t})
    elif op == "update":
        tid = args["task_id"]
        t = task_update(tid, user_id, title=args.get("title"), description=args.get("description"), priority=args.get("priority"), status=args.get("status"), due_date=args.get("due_date"), reminder_at=args.get("reminder_at"))
        if t:
            return json.dumps({"ok": True, "task": t})
        return json.dumps({"ok": False, "error": "Task not found"})
    elif op == "complete":
        tid = args["task_id"]
        t = task_complete(tid, user_id)
        if t:
            return json.dumps({"ok": True, "task": t})
        return json.dumps({"ok": False, "error": "Task not found"})
    elif op == "delete":
        task_delete(args["task_id"], user_id)
        return json.dumps({"ok": True})
    elif op == "list":
        tasks = task_list(user_id, args.get("status"))
        return json.dumps({"ok": True, "tasks": tasks})
    elif op == "get":
        t = task_get(args["task_id"], user_id)
        if t:
            return json.dumps({"ok": True, "task": t})
        return json.dumps({"ok": False, "error": "Task not found"})
    return json.dumps({"ok": False, "error": f"Unknown operation: {op}"})

_users_cache = None
_users_cache_time = 0


def load_users():
    global _users_cache, _users_cache_time
    now = time.time()
    if _users_cache is not None and now - _users_cache_time < 30:
        return _users_cache
    try:
        with open(USERS_FILE) as f:
            data = json.load(f)
        _users_cache = data.get("users", {})
        _users_cache_time = now
    except (FileNotFoundError, json.JSONDecodeError):
        _users_cache = {}
        _users_cache_time = now
    return _users_cache


def get_user_password(username):
    users = load_users()
    u = users.get(username)
    return u.get("password", "") if u else ""


def get_user_context_path(username):
    users = load_users()
    u = users.get(username)
    if u and u.get("context_file"):
        return os.path.join(u["context_file"])
    return ""


def read_user_context(username):
    path = get_user_context_path(username)
    print("Context path", path, "for", username)
    if path and os.path.exists(path):
        try:
            print("Reading", path)
            with open(path) as f:
                context = f.read()
                print(context)
                return context
        except:
            return ""
    return ""


def write_user_context(username, content):
    path = get_user_context_path(username)
    if path:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        existing = read_user_context(username)
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
        entry = f"[{timestamp}] {content}"
        new_content = (existing.strip() + "\n\n" + entry) if existing.strip() else entry
        with open(path, "w") as f:
            f.write(new_content)


_active_tokens = {}
_tokens_lock = threading.Lock()

_effective_contexts = {}
_effective_contexts_lock = threading.Lock()


def get_current_user(headers):
    token = headers.get("X-Auth-Token", "")
    with _tokens_lock:
        return _active_tokens.get(token)


sessions = {}
sessions_meta = {}
tasks = {}
model_status = "unloaded"
_last_tps = None
_last_llm_use = time.time()
_client_location = None

_data_lock = threading.Lock()

MAX_QUEUE_SIZE = 5
_task_queue = []
_queue_lock = threading.Lock()
_queue_cond = threading.Condition(_queue_lock)
_current_task_id = None

MAX_INPUT_TOKENS = 32768
AUTO_COMPACT_THRESHOLD = int(MAX_INPUT_TOKENS * 0.7)

_event_queue = _queue_mod.Queue()
_llm_pool = ThreadPoolExecutor(max_workers=1)
_tool_pool = ThreadPoolExecutor(max_workers=2)

_location_events = {}

_overheated = False
_gpu_temp = None
TEMP_THRESHOLD_ON = 85
TEMP_THRESHOLD_OFF = 65
RAM_EVAC_THRESHOLD = 95
RAM_RESUME_THRESHOLD = 70
_ram_evacuating = False


def load_sessions():
    global sessions, sessions_meta
    try:
        with open(SESSIONS_FILE) as f:
            data = json.load(f)
        with _data_lock:
            sessions = {}
            sessions_meta = {}
            for sid, sdata in data.get("sessions", {}).items():
                sessions[sid] = sdata.get("messages", [])
                sessions_meta[sid] = {
                    "name": sdata.get("name", "Chat"),
                    "created": sdata.get("created", time.time()),
                    "updated": sdata.get("updated", time.time()),
                    "user_id": sdata.get("user_id", ""),
                }
    except (FileNotFoundError, json.JSONDecodeError):
        with _data_lock:
            sessions = {}
            sessions_meta = {}


def save_sessions():
    with _data_lock:
        data = {"sessions": {}}
        for sid in sessions:
            meta = sessions_meta.get(
                sid, {"name": "Chat", "created": time.time(), "updated": time.time()}
            )
            data["sessions"][sid] = {
                "name": meta["name"],
                "created": meta["created"],
                "updated": meta["updated"],
                "user_id": meta.get("user_id", ""),
                "messages": sessions[sid],
            }
    with open(SESSIONS_FILE, "w") as f:
        json.dump(data, f, indent=2)


def _text_tokens(s):
    if not s:
        return 0
    # Multilingual/non-ASCII characters eat up far more tokens (~2 chars per token vs ~4 for English)
    non_ascii = sum(1 for ch in s if ord(ch) > 0x7F)
    divisor = 2.0 if non_ascii > len(s) * 0.15 else 4.0
    return int(len(s) / divisor)

def estimate_tokens(messages, include_tools=True):
    total = TOOLS_TOKEN_COST if include_tools else 0
    
    for msg in messages:
        total += PER_MESSAGE_OVERHEAD
        content = msg.get("content", "")
        
        # Standard string content
        if isinstance(content, str):
            total += _text_tokens(content)
            
        # Multi-modal array content (text + images + audio)
        elif isinstance(content, list):
            for part in content:
                if not isinstance(part, dict):
                    continue
                ptype = part.get("type")
                if ptype == "text":
                    total += _text_tokens(part.get("text", ""))
                elif ptype in ("image_url", "input_image", "image"):
                    total += IMAGE_TOKEN_COST
                elif ptype in ("audio_url", "input_audio", "audio"):
                    total += AUDIO_TOKEN_COST
                    
        # Tool call tokens
        for tc in msg.get("tool_calls") or []:
            total += _text_tokens(json.dumps(tc))

    # Return MUST be outside the for-loop!
    return max(1, total)


def trim_messages_for_context(messages):
    trimmed = list(messages)
    sys_msg = None
    if trimmed and trimmed[0].get("role") == "system":
        sys_msg = trimmed.pop(0)
    while estimate_tokens(trimmed) > MAX_INPUT_TOKENS and len(trimmed) > 1:
        trimmed.pop(0)
    if sys_msg:
        trimmed.insert(0, sys_msg)
    return trimmed


def compact_messages_copy(messages, keep_messages=6):
    """Return a compacted COPY of the message list (summary + recent messages)
    WITHOUT modifying the stored session. Old messages are summarized, not deleted."""
    msgs = list(messages)
    sys_msg = None
    if msgs and msgs[0].get("role") == "system":
        sys_msg = msgs.pop(0)
    if len(msgs) <= keep_messages + 1:
        return ([sys_msg] + msgs) if sys_msg else msgs
    to_compact = msgs[:-keep_messages] if keep_messages > 0 else msgs
    recent = msgs[-keep_messages:] if keep_messages > 0 else []
    compact_text = ""
    for m in to_compact:
        role = m.get("role", "unknown")
        content = m.get("content", "")
        if isinstance(content, list):
            parts = []
            for p in content:
                if isinstance(p, dict):
                    if p.get("type") == "text":
                        parts.append(p.get("text", ""))
            content = " ".join(parts)
        if not content:
            continue
        compact_text += f"[{role}]: {content}\n\n"
    if not compact_text.strip():
        return ([sys_msg] + msgs) if sys_msg else msgs
    summary = _summarize_with_llm(
        f"Compress the following conversation into a short paragraph, keeping all important details:\n\n{compact_text}"
    )
    if summary is None:
        return ([sys_msg] + msgs) if sys_msg else msgs
    new_msgs = []
    if sys_msg:
        new_msgs.append(sys_msg)
    new_msgs.append({"role": "system", "content": f"[Compressed context]: {summary}"})
    new_msgs.extend(recent)
    return new_msgs


def prepare_context_for_llm(sid, messages):
    """Build the message list to send to the LLM. When the conversation nears the
    context limit, old messages are summarized into a compressed context block —
    but the stored session is left untouched, so no messages are deleted."""
    total = estimate_tokens(messages)
    if total <= AUTO_COMPACT_THRESHOLD:
        context = trim_messages_for_context(messages)
        with _effective_contexts_lock:
            _effective_contexts.pop(sid, None)
        return context
    print(f"[context] Session {sid} estimate {total} tokens exceeds threshold {AUTO_COMPACT_THRESHOLD}; building compressed context for LLM")
    compacted = compact_messages_copy(messages)
    context = trim_messages_for_context(compacted)
    print(f"[context] Compressed context built; estimate after: {estimate_tokens(context)}")
    with _effective_contexts_lock:
        _effective_contexts[sid] = context
    return context


def effective_token_estimate(sid, messages):
    """Report the token count the UI shows: the compressed context actually sent
    to the LLM once compression has kicked in, falling back to the full history."""
    with _effective_contexts_lock:
        cached = _effective_contexts.get(sid)
    if cached is not None:
        return estimate_tokens(cached)
    return estimate_tokens(messages)


def context_token_report(sid, messages):
    """Token report for the UI: effective count sent to the LLM, the raw stored
    count, and whether context compression is currently active."""
    effective = effective_token_estimate(sid, messages)
    raw = estimate_tokens(messages)
    return {
        "token_estimate": effective,
        "raw_token_estimate": raw,
        "context_compressed": raw > effective,
    }


def _summarize_with_llm(text):
    payload = {
        "model": MODEL_ID,
        "messages": [
            {
                "role": "system",
                "content": "You summarize conversations concisely, preserving key facts, decisions, user preferences, and unresolved questions.",
            },
            {"role": "user", "content": text},
        ],
        "max_tokens": 1024,
        "temperature": 0.3,
        "stream": False,
    }
    try:
        r = requests.post(LLAMA_URL, json=payload, timeout=120)
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"[compact] LLM summarization failed: {e}")
        return None


def set_status(task_id, message):
    with _data_lock:
        if task_id in tasks:
            tasks[task_id]["status"] = "working"
            tasks[task_id]["message"] = message


def is_llama_alive():
    try:
        r = requests.get(f"{LLAMA_BASE}/health", timeout=5)
        return r.status_code == 200
    except:
        return False


_model_transition_lock = threading.Lock()


def unload_llama_model():
    global model_status
    with _model_transition_lock:
        with _data_lock:
            if model_status == "unloaded":
                return True
            model_status = "unloading"

        print("[llama] Requesting model unload from VRAM...")
        try:
            r = requests.post(
                f"{LLAMA_BASE}/models/unload", json={"model": MODEL_ID}, timeout=30
            )
            if r.status_code == 200:
                print("[llama] Model unloaded")
                with _data_lock:
                    model_status = "unloaded"
                return True
            print(f"[llama] Unload response: {r.status_code} {r.text[:200]}")
        except Exception as e:
            print(f"[llama] Unload error: {e}")

        # Check real status if unload failed or erred out
        with _data_lock:
            model_status = "chat_loaded" if is_llama_alive() else "unloaded"
        return False


def load_llama_model():
    global model_status, _last_llm_use
    with _data_lock:
        model_status = "loading"
    print(f"[llama] Sending load request for model '{MODEL_ID}'...")
    try:
        r = requests.post(
            f"{LLAMA_BASE}/models/load", json={"model": MODEL_ID}, timeout=180
        )
        if r.status_code in (200, 201):
            for i in range(30):
                if is_llama_alive():
                    print(f"[llama] Model ready (attempt {i+1})")
                    with _data_lock:
                        model_status = "chat_loaded"
                        _last_llm_use = time.time()  # Reset idle timer upon loading
                    return True
                time.sleep(2)
        else:
            print(f"[llama] Load failed ({r.status_code}): {r.text[:200]}")
    except Exception as e:
        print(f"[llama] Load exception: {e}")

    # Fallback check: verify if the server is alive and responding anyway
    if is_llama_alive():
        with _data_lock:
            model_status = "chat_loaded"
            _last_llm_use = time.time()  # Reset idle timer upon loading
        return True

    with _data_lock:
        model_status = "unloaded"
    return False


def _finalize_task(task_id, sid, msg_content, body):
    global _last_tps, _last_llm_use
    with _data_lock:
        t = tasks.get(task_id)
        if not t:
            return
        tools_used = list(t.get("_tools_used", []))
        search_details = list(t.get("_search_details", []))
        image_filename = t.get("image_file")
        gen_prompt = t.get("gen_prompt")
        image_model = t.get("_image_model")
    image_url = f"/output/{image_filename}" if image_filename else None
    if image_url:
        print(f"[finalize] image_file='{image_filename}' → image_url='{image_url}' for task {task_id}")  # DEBUG
    timings = body.get("timings", {})
    predicted_per_second = timings.get("predicted_per_second")
    reasoning = (
        body.get("choices", [{}])[0].get("message", {}).get("reasoning_content", "")
    )
    if not msg_content and reasoning:
        msg_content = "(No response content generated)"
    msg_entry = {
        "role": "assistant",
        "content": msg_content,
        "_reasoning": reasoning,
        "_tools_used": tools_used,
        "_image_url": image_url,
        "_gen_prompt": gen_prompt,
        "_image_model": image_model,
        "_search_details": search_details,
    }
    with _data_lock:
        if sid in sessions:
            sessions[sid].append(msg_entry)
            sessions_meta.setdefault(sid, {})["updated"] = time.time()
        _last_tps = predicted_per_second
        _last_llm_use = time.time()  # Reset idle timer when task finishes
    save_sessions()
    with _data_lock:
        if task_id in tasks:
            tasks[task_id] = {
                "status": "done",
                "response": msg_content,
                "session_id": sid,
                "session_name": sessions_meta.get(sid, {}).get("name", ""),
                **context_token_report(sid, sessions.get(sid, [])),
                "predicted_per_second": predicted_per_second,
                "tools_used": tools_used,
                "image": image_url,
                "_image_url": image_url,
                "gen_prompt": gen_prompt,
                "_image_model": image_model,
                "_search_details": search_details,
                "reasoning": reasoning,
            }


def free_comfyui_vram():
    print("[comfyui] Freeing VRAM...")
    try:
        r = requests.post(
            f"{COMFYUI_URL}/free",
            json={"unload_models": True, "free_memory": True},
            timeout=30,
        )
        if r.status_code == 200:
            print("[comfyui] VRAM freed")
            return True
    except Exception as e:
        print(f"[comfyui] Free error: {e}")
    finally:
        time.sleep(10)
    return False


def get_gpu_temp():
    try:
        r = subprocess.run(
            ["nvidia-smi", "--query-gpu=temperature.gpu", "--format=csv,noheader"],
            capture_output=True,
            text=True,
            timeout=5,
        )
        return int(r.stdout.strip())
    except Exception:
        return None


def get_ram_usage():
    try:
        r = subprocess.run(["free", "-m"], capture_output=True, text=True, timeout=5)
        lines = r.stdout.strip().split("\n")
        parts = lines[1].split()
        total = int(parts[1])
        available = int(parts[6])
        return (total - available) / total * 100
    except Exception:
        return None


def kill_llama_server():
    subprocess.run(["pkill", "-f", "llama-server"], capture_output=True)
    time.sleep(1)
    subprocess.run(["pkill", "-9", "-f", "llama-server"], capture_output=True)


def kill_comfyui():
    subprocess.run(["pkill", "-f", "main.py.*lowvram"], capture_output=True)


def restart_servers():
    print("Restarting servers")
    kill_llama_server()
    kill_comfyui()
    time.sleep(1)
    log_dir = os.path.expanduser("~/local-ai-files")
    llm_log = open(os.path.join(log_dir, "llama-server.log"), "a")
    comfy_log = open(os.path.join(log_dir, "comfyui.log"), "a")
    subprocess.Popen(
        [LLAMA_SERVER_PATH] + LLAMA_SERVER_ARGS,
        stdout=llm_log,
        stderr=llm_log,
        start_new_session=True,
    )
    subprocess.Popen(
        [
            os.path.join(VENV_PYTHON),
            "main.py",
            "--output-directory",
            COMFYUI_OUTPUT,
            "--input-directory",
            COMFYUI_INPUT,
            "--lowvram",
        ],
        cwd=COMFYUI_DIR,
        stdout=comfy_log,
        stderr=comfy_log,
        start_new_session=True,
    )
    deadline = time.time() + 120
    while time.time() < deadline:
        time.sleep(2)
        try:
            r = requests.get(f"{LLAMA_BASE}/health", timeout=3)
            if r.status_code == 200:
                print("[restart] llama-server healthy")
                return
        except Exception:
            pass
    print("[restart] llama-server did not respond within 2 minutes — killing")
    kill_llama_server()


def ensure_comfyui_running():
    try:
        r = requests.get(f"{COMFYUI_URL}/prompt", timeout=3)
        if r.status_code < 500:
            return
    except Exception:
        pass
    print("[comfyui] Not reachable — starting...")
    kill_comfyui()
    time.sleep(1)
    log_dir = os.path.expanduser("~/local-ai")
    comfy_log = open(os.path.join(log_dir, "comfyui.log"), "a")
    subprocess.Popen(
        [
            os.path.join(VENV_PYTHON),
            "main.py",
            "--output-directory",
            COMFYUI_OUTPUT,
            "--input-directory",
            COMFYUI_INPUT,
            "--lowvram",
        ],
        cwd=COMFYUI_DIR,
        stdout=comfy_log,
        stderr=comfy_log,
        start_new_session=True,
    )
    deadline = time.time() + 120
    while time.time() < deadline:
        time.sleep(2)
        try:
            r = requests.get(f"{COMFYUI_URL}/prompt", timeout=3)
            if r.status_code < 500:
                print("[comfyui] Healthy")
                return
        except Exception:
            pass
    print("[comfyui] Did not respond within 2 minutes")


def location_str():
    global _client_location
    if _client_location:
        return _client_location
    return None


def extract_city(loc):
    parts = [p.strip() for p in loc.split(",")]
    if len(parts) >= 3:
        return f"{parts[0]}, {parts[-3]}".strip(", ")
    if len(parts) >= 2:
        return parts[0].strip()
    return loc


def web_search(query, current_time=None, current_location=None):
    ts = datetime.now()
    from urllib.parse import urlencode

    clean_query = query.strip()
    params = {"q": clean_query, "format": "json"}
    search_url = f"{SEARXNG_URL}?{urlencode(params)}"
    print("Performing web search", search_url)
    try:
        r = requests.get(SEARXNG_URL, params=params, timeout=10)
        r.raise_for_status()
        print("Web-search completed")
        data = r.json()
    except Exception as e:
        print(f"Web-search failed: {e}")
        return json.dumps({
            "results": [],
            "search_date": ts.strftime("%Y-%m-%d %A"),
            "query": query,
            "search_url": search_url,
            "error": str(e),
        })
    results = data.get("results", [])[:5]
    formatted = []
    for x in results:
        formatted.append(
            {
                "title": x.get("title", ""),
                "url": x.get("url", ""),
                "snippet": x.get("content", "") or x.get("snippet", ""),
            }
        )
    return json.dumps(
        {
            "results": formatted,
            "search_date": ts.strftime("%Y-%m-%d %A"),
            "query": query,
            "search_url": search_url,
        }
    )


def fetch_page(url, max_chars=8000):
    import socket
    import ipaddress
    from bs4 import BeautifulSoup

    if not url:
        return json.dumps({"url": "", "error": "No URL provided."})
    try:
        parsed = urlparse(url)
        if parsed.scheme not in ("http", "https"):
            return json.dumps({"url": url, "error": "Only http/https URLs are supported."})
        host = parsed.hostname or ""
        ip = socket.gethostbyname(host)
        addr = ipaddress.ip_address(ip)
        if addr.is_private or addr.is_loopback or addr.is_link_local:
            return json.dumps({"url": url, "error": "Access to private/internal addresses is not allowed."})
    except Exception as e:
        return json.dumps({"url": url, "error": f"Invalid URL: {e}"})

    headers = {
        "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,text/plain;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
    }
    try:
        r = requests.get(url, headers=headers, timeout=15, allow_redirects=True)
        r.raise_for_status()
        ctype = r.headers.get("Content-Type", "").lower()
        if not any(t in ctype for t in ("text/html", "text/plain", "application/xhtml", "application/json", "application/xml")):
            return json.dumps({"url": url, "content_type": ctype, "error": "Skipped: page is not readable text content (likely binary/PDF/media)."})
        if not r.encoding:
            r.encoding = r.apparent_encoding
        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script", "style", "noscript", "svg", "nav", "footer", "header", "aside", "form"]):
            tag.decompose()
        title = soup.title.get_text(strip=True) if soup.title else ""
        main = soup.find("main") or soup.find("article") or soup.find("body") or soup
        text = main.get_text(separator="\n", strip=True)
        text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
        if len(text) > max_chars:
            text = text[:max_chars] + "\n...[truncated]"
        return json.dumps({
            "url": r.url,
            "title": title,
            "content": text or "(No readable text content extracted)",
        }, ensure_ascii=False)
    except Exception as e:
        print(f"[fetch_page] Failed: {e}")
        return json.dumps({"url": url, "error": f"Failed to fetch page: {e}"})


def generate_image(prompt, task_id, negative_prompt="", model="z_image"):
    global model_status
    print(f"\n[image] Generating image for task {task_id} with the prompt: {prompt}")
    set_status(task_id, "Freeing VRAM for image generation...")
    unload_llama_model()

    gen_tag = str(uuid.uuid4())[:8]
    prefix = f"gen_{gen_tag}_"
    cfg = IMAGE_MODELS.get(model, IMAGE_MODELS["z_image"])
    if model == "z_image":
        print("Chose Z-Image Turbo for image generation")
        workflow = {
            "62": {
                "class_type": "CLIPLoader",
                "inputs": {"clip_name": cfg["clip1"], "type": "lumina2"},
            },
            "63": {"class_type": "VAELoader", "inputs": {"vae_name": cfg["vae"]}},
            "66": {
                "class_type": "UNETLoader",
                "inputs": {"unet_name": cfg["unet"], "weight_dtype": "default"},
            },
            "67": {
                "class_type": "CLIPTextEncode",
                "inputs": {"text": prompt, "clip": ["62", 0]},
            },
            "68": {
                "class_type": "EmptySD3LatentImage",
                "inputs": {"width": 512, "height": 512, "batch_size": 1},
            },
            "69": {
                "class_type": "ModelSamplingAuraFlow",
                "inputs": {"shift": 3, "model": ["66", 0]},
            },
            "71": {
                "class_type": "CLIPTextEncode",
                "inputs": {"text": negative_prompt, "clip": ["62", 0]},
            },
            "70": {
                "class_type": "KSampler",
                "inputs": {
                    "seed": random.randint(0, 2**31),
                    "steps": 8,
                    "cfg": 1.0,
                    "sampler_name": "res_multistep",
                    "scheduler": "simple",
                    "denoise": 1.0,
                    "model": ["69", 0],
                    "positive": ["67", 0],
                    "negative": ["71", 0],
                    "latent_image": ["68", 0],
                },
            },
            "65": {
                "class_type": "VAEDecode",
                "inputs": {"samples": ["70", 0], "vae": ["63", 0]},
            },
            "9": {
                "class_type": "SaveImage",
                "inputs": {"filename_prefix": prefix, "images": ["65", 0]},
            },
        }
    elif model == "sd3_5_medium":
        print("Chose SD 3.5 for image generation")
        workflow = {
            "1": {"class_type": "UnetLoaderGGUF", "inputs": {"unet_name": cfg["unet"]}},
            "2": {
                "class_type": "TripleCLIPLoaderGGUF",
                "inputs": {
                    "clip_name1": cfg["clip1"],
                    "clip_name2": cfg["clip2"],
                    "clip_name3": cfg["t5"],
                    "type": "sd3",
                },
            },
            "3": {
                "class_type": "CLIPTextEncode",
                "inputs": {"text": prompt, "clip": ["2", 0]},
            },
            "4": {
                "class_type": "CLIPTextEncode",
                "inputs": {"text": negative_prompt, "clip": ["2", 0]},
            },
            "5": {
                "class_type": "EmptySD3LatentImage",
                "inputs": {"width": 512, "height": 512, "batch_size": 1},
            },
            "6": {
                "class_type": "KSampler",
                "inputs": {
                    "seed": random.randint(0, 2**31),
                    "steps": 20,  # Recommended steps for SD 3.5 Medium
                    "cfg": 4.5,  # Recommended CFG range for SD 3.5 Medium: 3.5 to 5.0
                    "sampler_name": "euler",
                    "scheduler": "sgm_uniform",
                    "denoise": 1.0,
                    "model": ["1", 0],
                    "positive": ["3", 0],
                    "negative": ["4", 0],
                    "latent_image": ["5", 0],
                },
            },
            "7": {"class_type": "VAELoader", "inputs": {"vae_name": cfg["vae"]}},
            "8": {
                "class_type": "VAEDecode",
                "inputs": {"samples": ["6", 0], "vae": ["7", 0]},
            },
            "9": {
                "class_type": "SaveImage",
                "inputs": {"filename_prefix": prefix, "images": ["8", 0]},
            },
        }
    else:
        print("No Image Model Selected Perfectly")

    with _data_lock:
        model_status = "image_active"
        tasks[task_id]["gen_prompt"] = prompt
        tasks[task_id]["_image_model"] = model
        tasks[task_id]["negative_prompt"] = negative_prompt
    ensure_comfyui_running()
    p_short = prompt[:200] + ("..." if len(prompt) > 200 else "")
    set_status(task_id, f"Generating image ({model})... Prompt: {p_short}")
    try:
        r = requests.post(
            f"{COMFYUI_URL}/prompt", json={"prompt": workflow}, timeout=120
        )
        data = r.json()

        if "error" in data:
            result = json.dumps({"error": f"ComfyUI: {data['error']}"})
        else:
            prompt_id = data["prompt_id"]
            found_file = None
            for _ in range(120):
                time.sleep(1)
                try:
                    hr = requests.get(f"{COMFYUI_URL}/history/{prompt_id}", timeout=10)
                    hist = hr.json()

                    if prompt_id in hist:
                        outputs = hist[prompt_id].get("outputs", {})
                        for node_id, node_out in outputs.items():
                            for img in node_out.get("images", []):
                                fname = img["filename"]
                                fpath = os.path.join(IMG_PATH, fname)
                                found_file = fpath
                                break
                        if found_file:
                            break
                except Exception:
                    pass
            if found_file:
                tasks[task_id]["image_file"] = found_file
                set_status(task_id, f"Image saved as {found_file}")
                print(f"[generate_image] SUCCESS: {found_file}")  # DEBUG
                result = json.dumps({"prompt_id": prompt_id, "file": found_file})
            else:
                print(f"[generate_image] TIMEOUT for task {task_id} after 120s")  # DEBUG
                result = json.dumps({"error": "Image generation timeout"})
    except Exception as e:
        result = json.dumps({"error": str(e)})
    finally:
        set_status(task_id, "Freeing image generation VRAM...")
        free_comfyui_vram()
        set_status(task_id, "Loading chat model...")
        load_llama_model()
    return result


def edit_image(
    prompt,
    task_id,
    image_b64,
    negative_prompt="",
    denoise=0.4,
    model="z_image",
    sid=None,
):
    print("Image edit called with denoise", denoise)
    if not image_b64 and sid:
        with _data_lock:
            msgs = list(sessions.get(sid, []))
        print(f"[edit_image] Scanning {len(msgs)} session messages for image sources")

        for msg in reversed(msgs):
            # 1. Check generated image URL attribute (_image_url)
            url = (msg.get("_image_url") or "").strip()
            if url:
                fname = os.path.basename(url)
                fpath = os.path.join(IMG_PATH, fname)
                print(
                    f"[edit_image] Checking _image_url path={fpath} exists={os.path.exists(fpath)}"
                )
                if os.path.exists(fpath):
                    with open(fpath, "rb") as f:
                        image_b64 = base64.b64encode(f.read()).decode()
                    break

            # 2. Check user-uploaded images stored in the message's content array
            content = msg.get("content")
            if isinstance(content, list):
                for part in reversed(content):
                    if isinstance(part, dict) and part.get("type") == "image_url":
                        img_url = part.get("image_url", {}).get("url", "")
                        if img_url.startswith("data:image"):
                            # Extracted base64 string directly from user upload
                            image_b64 = img_url.split(",", 1)[-1]
                            print(
                                "[edit_image] Extracted base64 image from user message content"
                            )
                            break
                if image_b64:
                    break

    if not image_b64:
        print("[edit_image] FAILED to find an image to edit")
        return json.dumps({"error": "No image provided for editing."})

    print(
        f"[edit_image] Found image ({len(image_b64)} bytes base64), proceeding with edit"
    )

    print(f"\n[image_edit] Editing image for task {task_id} with prompt: {prompt}")
    set_status(task_id, "Freeing VRAM for image editing...")
    unload_llama_model()

    gen_tag = str(uuid.uuid4())[:8]
    prefix = f"edit_{gen_tag}_"
    input_filename = f"input_{gen_tag}.png"

    input_dir = COMFYUI_INPUT
    os.makedirs(input_dir, exist_ok=True)
    input_filepath = os.path.join(input_dir, input_filename)

    with open(input_filepath, "wb") as f:
        f.write(base64.b64decode(image_b64))

    cfg = IMAGE_MODELS.get(model, IMAGE_MODELS["z_image"])

    workflow = {
        "62": {
            "class_type": "CLIPLoader",
            "inputs": {"clip_name": cfg["clip1"], "type": "lumina2"},
        },
        "63": {"class_type": "VAELoader", "inputs": {"vae_name": cfg["vae"]}},
        "66": {
            "class_type": "UNETLoader",
            "inputs": {"unet_name": cfg["unet"], "weight_dtype": "default"},
        },
        "67": {
            "class_type": "CLIPTextEncode",
            "inputs": {"text": prompt, "clip": ["62", 0]},
        },
        "71": {
            "class_type": "CLIPTextEncode",
            "inputs": {"text": negative_prompt, "clip": ["62", 0]},
        },
        "69": {
            "class_type": "ModelSamplingAuraFlow",
            "inputs": {"shift": 3, "model": ["66", 0]},
        },
        "5_load": {"class_type": "LoadImage", "inputs": {"image": input_filename}},
        "5_scale": {
            "class_type": "ImageScaleToTotalPixels",
            "inputs": {
                "image": ["5_load", 0],
                "megapixels": 0.262,  # ~512x512
                "upscale_method": "bicubic",
                "resolution_steps": 1,
            },
        },
        # Standard VAEEncode instead of VAEEncodeForInpaint
        "5_encode": {
            "class_type": "VAEEncode",
            "inputs": {"pixels": ["5_scale", 0], "vae": ["63", 0]},
        },
        "70": {
            "class_type": "KSampler",
            "inputs": {
                "seed": random.randint(0, 2**31),
                "steps": 8,
                "cfg": 1.0,
                "sampler_name": "res_multistep",
                "scheduler": "simple",
                "denoise": float(denoise),  # Dynamically controls edit depth
                "model": ["69", 0],
                "positive": ["67", 0],
                "negative": ["71", 0],
                "latent_image": ["5_encode", 0],
            },
        },
        "65": {
            "class_type": "VAEDecode",
            "inputs": {"samples": ["70", 0], "vae": ["63", 0]},
        },
        "9": {
            "class_type": "SaveImage",
            "inputs": {"filename_prefix": prefix, "images": ["65", 0]},
        },
    }
    with _data_lock:
        model_status = "image_active"
        tasks[task_id]["gen_prompt"] = prompt
        tasks[task_id]["_image_model"] = model
        tasks[task_id]["negative_prompt"] = negative_prompt

    ensure_comfyui_running()
    set_status(task_id, f"Editing image ({model})... Prompt: {prompt[:150]}")

    try:
        r = requests.post(
            f"{COMFYUI_URL}/prompt", json={"prompt": workflow}, timeout=120
        )
        data = r.json()

        if "error" in data:
            result = json.dumps({"error": f"ComfyUI: {data['error']}"})
        else:
            prompt_id = data["prompt_id"]
            found_file = None
            for _ in range(120):
                time.sleep(1)
                try:
                    hr = requests.get(f"{COMFYUI_URL}/history/{prompt_id}", timeout=10)
                    hist = hr.json()
                    if prompt_id in hist:
                        outputs = hist[prompt_id].get("outputs", {})
                        for node_id, node_out in outputs.items():
                            for img in node_out.get("images", []):
                                fname = img["filename"]
                                found_file = os.path.join(IMG_PATH, fname)
                                break
                        if found_file:
                            break
                except Exception:
                    pass

            if found_file:
                tasks[task_id]["image_file"] = found_file
                set_status(task_id, f"Edited image saved as {found_file}")
                result = json.dumps({"prompt_id": prompt_id, "file": found_file})
            else:
                result = json.dumps({"error": "Image editing timeout"})
    except Exception as e:
        result = json.dumps({"error": str(e)})
    finally:
        if os.path.exists(input_filepath):
            try:
                os.remove(input_filepath)
                print(f"[edit_image] Cleaned up input file: {input_filepath}")
            except Exception as e:
                print(f"[edit_image] Failed to cleanup input file: {e}")
        set_status(task_id, "Freeing image generation VRAM...")
        free_comfyui_vram()
        set_status(task_id, "Loading chat model...")
        load_llama_model()

    return result


def _event_post(ev_type, task_id, **data):
    _event_queue.put((ev_type, task_id, data))


def _llm_worker(task_id, sid, round_num, msgs):
    try:
        if estimate_tokens(msgs) > AUTO_COMPACT_THRESHOLD:
            set_status(task_id, "Context is full — compressing older messages...")
        messages = prepare_context_for_llm(sid, msgs)
        tool_msgs = [m for m in messages if isinstance(m, dict) and m.get("role") == "tool"]
        if tool_msgs:
            print(f"[llm_round] Round {round_num} includes {len(tool_msgs)} tool message(s) with search results")  # DEBUG
        payload = {
            "model": MODEL_ID,
            "messages": messages,
            "tools": TOOLS,
            "tool_choice": "auto",
            "max_tokens": MAX_INPUT_TOKENS,
            #"reasoning_budget": REASONING_BUDGET,
            #"reasoning_effort": "medium",
        }
        has_image = any(
            isinstance(m.get("content"), list)
            and any(p.get("type") == "image_url" for p in m["content"])
            for m in payload.get("messages", [])
        )
        payload["stream"] = True
        r = requests.post(LLAMA_URL, json=payload, stream=True, timeout=600)
        r.encoding = "utf-8"
        reasoning_buf = ""
        content_buf = ""
        tool_calls_map = {}
        with _data_lock:
            prev_reasoning = tasks.get(task_id, {}).get("reasoning", "")
        for line in r.iter_lines(decode_unicode=True):
            if not line or not line.startswith("data: "):
                continue
            data_str = line[6:]
            if data_str.strip() == "[DONE]":
                break
            try:
                chunk = json.loads(data_str)
            except json.JSONDecodeError:
                continue
            choices = chunk.get("choices", [])
            if not choices:
                continue
            delta = choices[0].get("delta", {})
            rc = delta.get("reasoning_content")
            if rc:
                reasoning_buf += rc
                with _data_lock:
                    if task_id in tasks:
                        tasks[task_id]["reasoning"] = prev_reasoning + reasoning_buf
            c = delta.get("content")
            if c:
                content_buf += c
            tc_list = delta.get("tool_calls")
            if tc_list:
                for tc in tc_list:
                    idx = tc.get("index", 0)
                    if idx not in tool_calls_map:
                        fn = tc.get("function", {})
                        tool_calls_map[idx] = {
                            "index": idx,
                            "id": tc.get("id", ""),
                            "type": tc.get("type", "function"),
                            "function": {
                                "name": fn.get("name", ""),
                                "arguments": fn.get("arguments", ""),
                            },
                        }
                    else:
                        existing = tool_calls_map[idx]
                        if tc.get("id"):
                            existing["id"] = tc["id"]
                        fn = tc.get("function")
                        if fn:
                            if fn.get("name"):
                                existing["function"]["name"] = fn["name"]
                            if fn.get("arguments"):
                                existing["function"]["arguments"] += fn["arguments"]
        print(f"[llm_round] Round {round_num} done: reasoning_buf={len(reasoning_buf)} chars, content_buf={len(content_buf)} chars, tool_calls={len(tool_calls_map)}")  # DEBUG
        msg = {
            "role": "assistant",
            "content": content_buf,
            "reasoning_content": prev_reasoning + reasoning_buf,
        }
        if tool_calls_map:
            msg["tool_calls"] = list(tool_calls_map.values())
        body = {"choices": [{"message": msg}]}
        if "choices" in body:
            _event_post("llm_ok", task_id, body=body, round=round_num, sid=sid)
        else:
            _event_post(
                "llm_err",
                task_id,
                error="Unexpected response",
                round=round_num,
                sid=sid,
            )
    except Exception as e:
        err_text = str(e)
        if "image" in err_text.lower() or "vision" in err_text.lower():
            err_text = "The current model does not support image input. Please use a vision-capable model or send text-only messages."
        _event_post("llm_err", task_id, error=err_text, round=round_num, sid=sid)


def _tool_worker(task_id, sid, tc, image_b64, round_num, tool_index):
    tool_name = tc["function"]["name"]
    try:
        args = json.loads(tc["function"]["arguments"])
    except Exception:
        args = {}

    with _data_lock:
        tu = list(tasks.get(task_id, {}).get("_tools_used", []))
    has_generated_image = "generate_image" in tu

    if tool_name == "get_user_location":
        if _client_location:
            result = _client_location
        else:
            ev = threading.Event()
            _location_events[task_id] = ev
            set_status(task_id, "location_needed")
            ev.wait(timeout=60)
            _location_events.pop(task_id, None)
            result = _client_location if _client_location else "User denied location access"
        _event_post("tool_ok", task_id, tc_id=tc["id"], result=result, sid=sid, round=round_num, tool_index=tool_index)
        return

    if tool_name == "read_file":
        file_url = args.get("file_url", "")
        filename = os.path.basename(urlparse(file_url).path)
        fpath = os.path.abspath(os.path.join(UPLOADS_DIR, filename))
        if fpath.startswith(os.path.abspath(UPLOADS_DIR)) and os.path.exists(fpath):
            text = read_file_text(fpath)
            if text:
                result = f"Content of {file_url}:\n\n{text}"
            else:
                result = f"Could not extract text from {file_url}. The file may contain only images."
        else:
            result = f"File not found: {file_url}"
        _event_post("tool_ok", task_id, tc_id=tc["id"], result=result, sid=sid, round=round_num, tool_index=tool_index)
        return

    if tool_name == "web_search":
        set_status(task_id, f"Searching web for: {args.get('query')}...")
        with _data_lock:
            client_ts = tasks.get(task_id, {}).get("_client_timestamp")
        try:
            result = web_search(
                args["query"],
                current_time=args.get("current_time"),
                current_location=args.get("current_location"),
            )
        except Exception as e:
            print(f"[web_search] Unhandled exception for task {task_id}: {e}")
            result = json.dumps({"results": [], "query": args.get("query"), "error": str(e)})
        print(f"[web_search] RAW result for task {task_id}: {result[:300]}...")  # DEBUG
        with _data_lock:
            t = tasks.get(task_id)
            if t:
                t.setdefault("_tools_used", []).append(tool_name)
                try:
                    t.setdefault("_search_details", []).append(json.loads(result))
                except Exception:
                    pass
        llm_result = (
            f"Web search results for query '{args.get('query')}'. "
            f"Analyze these search results thoroughly and provide a clear, accurate response based on the findings:\n\n{result}"
        )
        print(f"[web_search] LLM-bound result (with analysis instruction) for task {task_id}: {llm_result[:400]}...")  # DEBUG
        _event_post(
            "tool_ok",
            task_id,
            tc_id=tc["id"],
            result=llm_result,
            sid=sid,
            round=round_num,
            tool_index=tool_index,
        )

    elif tool_name == "fetch_page":
        set_status(task_id, f"Fetching page: {args.get('url', '')}...")
        try:
            result = fetch_page(args.get("url", ""))
        except Exception as e:
            print(f"[fetch_page] Unhandled exception for task {task_id}: {e}")
            result = json.dumps({"url": args.get("url", ""), "error": str(e)})
        print(f"[fetch_page] Result for task {task_id}: {result[:300]}...")  # DEBUG
        with _data_lock:
            t = tasks.get(task_id)
            if t:
                t.setdefault("_tools_used", []).append(tool_name)
                try:
                    res = json.loads(result)
                    t.setdefault("_search_details", []).append({
                        "tool": "fetch_page",
                        "url": res.get("url", args.get("url", "")),
                        "title": res.get("title", ""),
                        "content": res.get("content", ""),
                        "error": res.get("error", ""),
                    })
                except Exception:
                    pass
        llm_result = (
            f"Page content fetched from URL '{args.get('url')}'. "
            f"Use this content to answer the user's question accurately. "
            f"If the content is insufficient or was truncated, you may fetch another page or fall back to the search results:\n\n{result}"
        )
        _event_post(
            "tool_ok",
            task_id,
            tc_id=tc["id"],
            result=llm_result,
            sid=sid,
            round=round_num,
            tool_index=tool_index,
        )

    elif tool_name == "edit_image":
        result = edit_image(
            prompt=args.get("prompt", ""),
            task_id=task_id,
            image_b64=image_b64,
            negative_prompt=args.get("negative_prompt", ""),
            denoise=args.get("denoise", 0.4),
            model="z_image",
            sid=sid,
        )
        res_data = json.loads(result)
        if "file" in res_data:
            fn = os.path.basename(res_data['file'])
            image_url = f"/output/{fn}"
            with _data_lock:
                t = tasks.get(task_id)
                if t:
                    t.setdefault("_tools_used", []).append(tool_name)
                    t["image_file"] = fn
                    t["gen_prompt"] = args.get("prompt", "")
                    t["_image_model"] = None
            tool_result = json.dumps({
                "image_url": image_url,
                "prompt": args.get("prompt", ""),
                "model": None,
            })
            _event_post(
                "tool_ok",
                task_id,
                tc_id=tc["id"],
                result=tool_result,
                sid=sid,
                round=round_num,
                tool_index=tool_index,
            )
        else:
            _event_post(
                "tool_ok",
                task_id,
                tc_id=tc["id"],
                result=result,
                sid=sid,
                round=round_num,
                tool_index=tool_index,
            )

    elif tool_name == "generate_image":
        if has_generated_image:
            result = json.dumps(
                {"error": "Image generation limit reached for this prompt."}
            )
            _event_post(
                "tool_ok",
                task_id,
                tc_id=tc["id"],
                result=result,
                sid=sid,
                round=round_num,
                tool_index=tool_index,
            )
        else:
            result = generate_image(
                prompt=args.get("prompt", ""),
                task_id=task_id,
                negative_prompt=args.get("negative_prompt", ""),
                model=args.get("model") or "z_image",
            )
            res_data = json.loads(result)
            if "file" in res_data:
                fn = os.path.basename(res_data['file'])
                image_url = f"/output/{fn}"
                image_model_s = args.get("model") or "z_image"
                print(f"[tool_worker] Image file: {res_data['file']}, basename: {fn}, url: {image_url}")  # DEBUG
                with _data_lock:
                    t = tasks.get(task_id)
                    if t:
                        t.setdefault("_tools_used", []).append(tool_name)
                        t["image_file"] = fn
                        t["gen_prompt"] = args.get("prompt", "")
                        t["_image_model"] = image_model_s
                        print(f"[tool_worker] Stored image_file='{fn}' in task {task_id}")  # DEBUG
                tool_result = json.dumps({
                    "image_url": image_url,
                    "prompt": args.get("prompt", ""),
                    "model": image_model_s,
                })
                _event_post(
                    "tool_ok",
                    task_id,
                    tc_id=tc["id"],
                    result=tool_result,
                    sid=sid,
                    round=round_num,
                    tool_index=tool_index,
                )
            else:
                print(f"[tool_worker] generate_image FAILED for task {task_id}: {result[:200]}")  # DEBUG
                _event_post(
                    "tool_ok",
                    task_id,
                    tc_id=tc["id"],
                    result=result,
                    sid=sid,
                    round=round_num,
                    tool_index=tool_index,
                )
    elif tool_name == "update_user_context":
        content = args.get("content", "")
        user = ""
        with _data_lock:
            t = tasks.get(task_id)
            if t:
                user = t.get("_user", "")
        if user:
            write_user_context(user, content)
            print(f"[context] Updated context for user '{user}' ({len(content)} chars)")
        result = json.dumps({"status": "ok", "saved": bool(user)})
        _event_post(
            "tool_ok",
            task_id,
            tc_id=tc["id"],
            result=result,
            sid=sid,
            round=round_num,
            tool_index=tool_index,
        )
    elif tool_name == "manage_tasks":
        user = ""
        with _data_lock:
            t = tasks.get(task_id)
            if t:
                user = t.get("_user", "")
        if not user:
            result = json.dumps({"ok": False, "error": "User not found"})
        else:
            result = handle_task_tool(user, args)
        _event_post(
            "tool_ok",
            task_id,
            tc_id=tc["id"],
            result=result,
            sid=sid,
            round=round_num,
            tool_index=tool_index,
        )
    else:
        result = json.dumps({"error": f"Unknown tool: {tool_name}"})
        _event_post(
            "tool_ok",
            task_id,
            tc_id=tc["id"],
            result=result,
            sid=sid,
            round=round_num,
            tool_index=tool_index,
        )


def _prepare_session(task_id, sid, user_message, image_b64, audio_b64=None, client_ts=None):
    try:
        if client_ts:
            ts = datetime.fromisoformat(client_ts.replace("Z", "+00:00"))
        else:
            ts = datetime.now()
    except Exception:
        ts = datetime.now()
    loc = location_str()
    loc_context = f" [User location: {loc}]" if loc else ""
    date_loc_context = f"[Current date: {ts.strftime('%Y-%m-%d %A %H:%M')}]{loc_context}"
    user = ""
    with _data_lock:
        t = tasks.get(task_id)
        if t:
            user = t.get("_user", "")
    user_context = read_user_context(user) if user else ""
    context_block = f"\n\n## User Context\n{user_context}" if user_context else ""
    full_sys_content = f"{SYS_CONTENT}\n\n{date_loc_context}{context_block}"
    full_sys_content = full_sys_content.replace(
        "%current_time%", ts.strftime("%Y-%m-%d %A %H:%M")
    )
    if loc:
        full_sys_content = full_sys_content.replace("%current_location%", loc)
    else:
        full_sys_content = full_sys_content.replace("Currently the server is hosted on %current_location%.", "")
        full_sys_content = full_sys_content.replace("%current_location%", "not available")
    if user_context:
        print(
            f"[context] Injected {len(user_context)} chars of context for user '{user}'"
        )
    with _data_lock:
        if sid not in sessions or not sessions[sid]:
            sessions[sid] = [{"role": "system", "content": full_sys_content}]
        elif sessions[sid][0].get("role") == "system":
            sessions[sid][0]["content"] = full_sys_content
        else:
            sessions[sid].insert(0, {"role": "system", "content": full_sys_content})
        if sid not in sessions_meta:
            sessions_meta[sid] = {
                "name": user_message[:50],
                "created": time.time(),
                "updated": time.time(),
            }
        content = []
        if image_b64:
            content.append(
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/jpeg;base64,{image_b64}"},
                }
            )
        if audio_b64:
            content.append(
                {
                    "type": "audio_url",
                    "audio_url": {"url": f"data:audio/webm;base64,{audio_b64}"},
                }
            )
        content.append(
            {
                "type": "text",
                "text": user_message,
            }
        )
        sessions[sid].append({"role": "user", "content": content, "_timestamp": datetime.now().isoformat()})
        if sessions_meta[sid]["name"] in ("New Chat", ""):
            sessions_meta[sid]["name"] = user_message[:50] + (
                "..." if len(user_message) > 50 else ""
            )
        sessions_meta[sid]["updated"] = time.time()
    save_sessions()
    with _data_lock:
        ms = model_status
    if ms != "chat_loaded":
        load_llama_model()


def _start_llm_round(task_id, sid, round_num):
    with _data_lock:
        ms = model_status
    if ms != "chat_loaded":
        load_llama_model()
    with _data_lock:
        t = tasks.get(task_id)
        if not t:
            return
        t["_state"] = "llm_waiting"
        t["_round"] = round_num
        messages = list(sessions.get(sid, []))
    print(f"[llm_round] Starting round {round_num} for task {task_id} with {len(messages)} raw messages")  # DEBUG
    set_status(
        task_id, "Thinking..." if round_num == 0 else f"Thinking (round {round_num})..."
    )
    _llm_pool.submit(_llm_worker, task_id, sid, round_num, messages)


def _set_task_error(task_id, error, sid=None):
    with _data_lock:
        if task_id in tasks:
            d = tasks[task_id]
            tasks[task_id] = {
                "status": "error",
                "error": str(error),
                "session_id": d.get("session_id", sid),
            }


def _event_loop():
    global _current_task_id
    while True:
        ev_type, task_id, data = _event_queue.get()
        t = tasks.get(task_id)
        if not t:
            continue

        if ev_type == "start":
            sid = data["sid"]
            user_message = data["message"]
            image_b64 = data.get("image")
            audio_b64 = data.get("audio")
            user = data.get("user", "")
            client_ts = data.get("client_timestamp")
            with _data_lock:
                tasks[task_id] = {
                    "status": "working",
                    "message": "Processing task...",
                    "session_id": sid,
                    "_tools_used": [],
                    "_search_details": [],
                    "_original_message": user_message,
                    "_original_image": image_b64,
                    "_audio": audio_b64,
                    "_user": user,
                    "_client_timestamp": client_ts,
                }
            _current_task_id = task_id
            _prepare_session(task_id, sid, user_message, image_b64, audio_b64, client_ts)
            _start_llm_round(task_id, sid, 0)

        elif ev_type == "llm_ok":
            if t.get("_state") != "llm_waiting":
                continue
            sid = data["sid"]
            round_num = data["round"]
            body = data["body"]
            msg = body["choices"][0]["message"]
            with _data_lock:
                _last_llm_use = time.time()
            if msg.get("tool_calls"):
                with _data_lock:
                    tt = tasks.get(task_id)
                    if tt:
                        tt.setdefault("_tools_used", [])
                        tt.setdefault("_search_details", [])
                        rc = msg.get("reasoning_content", "")
                        if rc:
                            tt["reasoning"] = rc
                pending = len(msg["tool_calls"])
                print(f"[llm_ok] Round {round_num}: LLM requested {pending} tool(s) for task {task_id}")  # DEBUG
                with _data_lock:
                    tt = tasks.get(task_id)
                    if tt:
                        tt["_state"] = "tools_running"
                        tt["_pending_tools"] = pending
                with _data_lock:
                    if sid in sessions:
                        assistant_msg = {"role": "assistant"}
                        if msg.get("content"):
                            assistant_msg["content"] = msg["content"]
                        if msg.get("tool_calls"):
                            assistant_msg["tool_calls"] = msg["tool_calls"]
                        sessions[sid].append(assistant_msg)
                        sessions_meta.setdefault(sid, {})["updated"] = time.time()
                save_sessions()
                for i, tc in enumerate(msg["tool_calls"]):
                    _tool_pool.submit(
                        _tool_worker,
                        task_id,
                        sid,
                        tc,
                        t.get("_original_image"),
                        round_num,
                        i,
                    )
            else:
                print(f"[llm_ok] Round {round_num}: LLM generated final response (no tool calls) for task {task_id}")  # DEBUG
                _finalize_task(task_id, sid, (msg.get("content") or ""), body)

        elif ev_type == "llm_err":
            if t.get("_state") != "llm_waiting":
                continue
            _set_task_error(task_id, data["error"], data.get("sid"))

        elif ev_type == "tool_ok":
            sid = data["sid"]
            tc_id = data["tc_id"]
            result = data["result"]
            with _data_lock:
                if sid in sessions:
                    sessions[sid].append(
                        {"role": "tool", "tool_call_id": tc_id, "content": result}
                    )
                    sessions_meta.setdefault(sid, {})["updated"] = time.time()
                    print(f"[tool_ok] Appended tool result to session {sid} for task {task_id}")  # DEBUG
                tt = tasks.get(task_id)
                if not tt or tt.get("status") in ("done", "error"):
                    continue
                pending = (tt.get("_pending_tools", 0) - 1) if tt else 0
                if tt:
                    tt["_pending_tools"] = pending
            save_sessions()
            print(f"[tool_ok] Pending tools left for task {task_id}: {pending}")  # DEBUG
            if pending <= 0:
                round_num = data.get("round", 0) + 1
                print(f"[tool_ok] All tools done for task {task_id}. Starting LLM round {round_num} with search results in context.")  # DEBUG
                with _data_lock:
                    tt = tasks.get(task_id)
                    if tt:
                        tt["_round"] = round_num
                if round_num < 10:
                    _start_llm_round(task_id, sid, round_num)
                else:
                    _set_task_error(task_id, "Max tool rounds exceeded", sid)

        elif ev_type == "tool_err":
            result = data.get(
                "result", json.dumps({"error": data.get("error", "Tool error")})
            )
            with _data_lock:
                if data.get("sid") in sessions:
                    sessions[data["sid"]].append(
                        {
                            "role": "tool",
                            "tool_call_id": data["tc_id"],
                            "content": result,
                        }
                    )
                    sessions_meta.setdefault(data["sid"], {})["updated"] = time.time()
                tt = tasks.get(task_id)
                if not tt or tt.get("status") in ("done", "error"):
                    continue
                pending = (tt.get("_pending_tools", 0) - 1) if tt else 0
                if tt:
                    tt["_pending_tools"] = pending
            save_sessions()
            if pending <= 0:
                round_num = data.get("round", 0) + 1
                with _data_lock:
                    tt = tasks.get(task_id)
                    if tt:
                        tt["_round"] = round_num
                if round_num < 10:
                    _start_llm_round(task_id, data["sid"], round_num)
                else:
                    _set_task_error(task_id, "Max tool rounds exceeded", data["sid"])


def _queue_worker():
    global _current_task_id
    while True:
        item = None
        with _queue_lock:
            while not _task_queue:
                _queue_cond.wait()
            with _data_lock:
                oh = _overheated
            if oh or _ram_evacuating:
                label = "GPU overheating" if oh else "RAM pressure — restarting servers"
                for qitem in _task_queue:
                    tid = qitem["task_id"]
                    if tid in tasks:
                        tasks[tid] = {
                            "status": "waiting",
                            "message": f"Server paused — {label}. Will resume shortly.",
                            "session_id": qitem["session_id"],
                        }
                _queue_cond.wait(5)
                continue
            item = _task_queue.pop(0)
            _current_task_id = item["task_id"]
        _event_post(
            "start",
            item["task_id"],
            sid=item["session_id"],
            message=item["message"],
            image=item.get("image"),
            audio=item.get("audio"),
            user=item.get("user", ""),
            client_timestamp=item.get("client_timestamp"),
        )
        # Wait for this task to finish (status becomes "done" or "error") before dequeuing the next
        while True:
            with _data_lock:
                st = tasks.get(item["task_id"], {}).get("status")
            if st in ("done", "error"):
                break
            time.sleep(0.5)
        with _queue_lock:
            _current_task_id = None
            _queue_cond.notify_all()


def _idle_unload_loop():
    global _last_llm_use
    while True:
        time.sleep(10)

        with _queue_lock:
            queue_active = len(_task_queue) > 0 or _current_task_id is not None

        with _data_lock:
            ms = model_status
            lu = _last_llm_use

        # Only unload if loaded, inactive for > 300s, and no queue tasks pending
        if ms == "chat_loaded" and (time.time() - lu > 300) and not queue_active:
            print("[idle] No LLM activity for 300s, releasing VRAM model weights...")
            unload_llama_model()


def _reminder_loop():
    while True:
        try:
            now = datetime.now().isoformat()
            due = _db_fetch("SELECT * FROM tasks WHERE reminder_at IS NOT NULL AND reminder_at <= ? AND reminded=0 AND status NOT IN ('completed','cancelled')", (now,))
            for task in due:
                print(f"[reminder] Task '{task['title']}'. User: {task['user_id']}")
                _db_run("UPDATE tasks SET reminded=1 WHERE id=?", (task["id"],))
        except Exception as e:
            print(f"[reminder] Error: {e}")
        time.sleep(30)


def _evacuate_ram():
    global _current_task_id, _ram_evacuating
    _ram_evacuating = True
    print("[ram] Emergency RAM evacuation")
    with _queue_lock:
        tid = _current_task_id
        if tid:
            with _data_lock:
                t = tasks.get(tid)
                if t and t.get("status") not in ("done", "error"):
                    entry = {
                        "task_id": tid,
                        "session_id": t.get("session_id", ""),
                        "message": t.get("_original_message", ""),
                        "image": t.get("_original_image"),
                    }
                    _task_queue.insert(0, entry)
                    t["status"] = "error"
                    t["error"] = "Server ran out of RAM — requeued"
                    t["_ram_evacuating"] = True
                    print(f"[ram] Requeued task {tid} to front of queue")
    kill_llama_server()
    kill_comfyui()
    print("[ram] Killed llama-server and ComfyUI")
    while True:
        time.sleep(5)
        ram = get_ram_usage()
        if ram is not None and ram <= RAM_RESUME_THRESHOLD:
            print(f"[ram] RAM {ram:.0f}% ≤ {RAM_RESUME_THRESHOLD}%, restarting servers")
            break
    restart_servers()
    _ram_evacuating = False


def _thermal_monitor():
    global _overheated, _gpu_temp
    while True:
        time.sleep(10)
        temp = get_gpu_temp()
        with _data_lock:
            _gpu_temp = temp
            if temp is not None and temp >= TEMP_THRESHOLD_ON:
                if not _overheated:
                    print(
                        f"[thermal] GPU {temp}°C >= {TEMP_THRESHOLD_ON}°C, OVERHEATED"
                    )
                    _overheated = True
            elif _overheated and (temp is None or temp <= TEMP_THRESHOLD_OFF):
                print(f"[thermal] GPU {temp}°C <= {TEMP_THRESHOLD_OFF}°C, resumed")
                _overheated = False

        if _overheated:
            with _queue_lock:
                busy = _current_task_id is not None
            if not busy:
                with _data_lock:
                    ms = model_status
                if ms == "chat_loaded":
                    print("[thermal] Overheated — unloading chat model")
                    unload_llama_model()
                elif ms == "image_active":
                    print("[thermal] Overheated — freeing ComfyUI VRAM")
                    free_comfyui_vram()

        if not _ram_evacuating:
            ram = get_ram_usage()
            if ram is not None and ram >= RAM_EVAC_THRESHOLD:
                print(f"[ram] RAM usage {ram:.0f}% >= {RAM_EVAC_THRESHOLD}%")
                _evacuate_ram()


def read_index_html():
    p = os.path.join(os.path.dirname(os.path.abspath(__file__)), "dist", "index.html")
    try:
        with open(p) as f:
            return f.read()
    except:
        return "<html><body><h1>index.html missing</h1></body></html>"


def strip_html(text):
    import re
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL)
    text = re.sub(r"<script[^>]*>.*?</script>", "", text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def read_file_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    with open(file_path, "rb") as f:
        raw = f.read()
    if ext == ".pdf":
        try:
            import fitz
            doc = fitz.open(stream=raw, filetype="pdf")
            lines = []
            for page in doc:
                lines.append(page.get_text())
            doc.close()
            text = "\n".join(lines)
            return strip_html(text)
        except ImportError:
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmpf:
                tmpf.write(raw)
                tmp = tmpf.name
            try:
                r = subprocess.run(
                    ["pdftotext", tmp, "-"], capture_output=True, text=True, timeout=30
                )
                return r.stdout
            finally:
                os.unlink(tmp)
    elif ext == ".docx":
        from docx import Document
        doc = Document(io.BytesIO(raw))
        return "\n".join(p.text for p in doc.paragraphs)
    elif ext == ".doc":
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmpf:
            tmpf.write(raw)
            tmp = tmpf.name
        try:
            r = subprocess.run(
                ["catdoc", tmp], capture_output=True, text=True, timeout=30
            )
            if r.returncode == 0:
                return r.stdout
            r = subprocess.run(
                ["antiword", tmp], capture_output=True, text=True, timeout=30
            )
            return r.stdout
        finally:
            os.unlink(tmp)
    elif ext in (".xls", ".xlsx"):
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        rows = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                rows.append("\t".join(str(c) if c is not None else "" for c in row))
        wb.close()
        return "\n".join(rows)
    return ""


def extract_file_text(name, data_b64):
    ext = os.path.splitext(name)[1].lower()
    raw = base64.b64decode(data_b64)
    if ext == ".pdf":
        try:
            import fitz

            doc = fitz.open(stream=raw, filetype="pdf")
            lines = [page.get_text() for page in doc]
            doc.close()
            text = "\n".join(lines)
            return strip_html(text)
        except ImportError:
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                f.write(raw)
                tmp = f.name
            try:
                r = subprocess.run(
                    ["pdftotext", tmp, "-"], capture_output=True, text=True, timeout=30
                )
                return r.stdout
            finally:
                os.unlink(tmp)
    elif ext == ".docx":
        from docx import Document

        doc = Document(io.BytesIO(raw))
        return "\n".join(p.text for p in doc.paragraphs)
    elif ext == ".doc":
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as f:
            f.write(raw)
            tmp = f.name
        try:
            r = subprocess.run(
                ["catdoc", tmp], capture_output=True, text=True, timeout=30
            )
            if r.returncode == 0:
                return r.stdout
            r = subprocess.run(
                ["antiword", tmp], capture_output=True, text=True, timeout=30
            )
            return r.stdout
        finally:
            os.unlink(tmp)
    elif ext in (".xls", ".xlsx"):
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        rows = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                rows.append("\t".join(str(c) if c is not None else "" for c in row))
        wb.close()
        return "\n".join(rows)
    return ""


class Handler(http.server.SimpleHTTPRequestHandler):
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header(
            "Access-Control-Allow-Methods", "GET, POST, PUT, DELETE, OPTIONS"
        )
        self.send_header("Access-Control-Allow-Headers", "Content-Type, X-Auth-Token")
        self.end_headers()

    def do_GET(self):
        if self.path == "/api/user-context":
            user = get_current_user(self.headers)
            if not user:
                self.send_json({"error": "Unauthorized"}, status=401)
                return
            context = read_user_context(user)
            self.send_json(
                {
                    "context": context,
                    "username": user,
                    "context_file": get_user_context_path(user),
                }
            )
        elif self.path == "/api/check-auth":
            user = get_current_user(self.headers)
            if user:
                self.send_json({"authenticated": True, "username": user})
            else:
                self.send_json({"authenticated": False})
        elif self.path == "/api/model-status":
            with _data_lock:
                ms, tps, oh, gtemp = model_status, _last_tps, _overheated, _gpu_temp
            try:
                user = get_current_user(self.headers)
                reminder_count = len(_db_fetch("SELECT id FROM tasks WHERE user_id=? AND reminder_at IS NOT NULL AND reminder_at <= ? AND reminded=0 AND status NOT IN ('completed','cancelled')", (user, datetime.now().isoformat()))) if user else 0
            except Exception:
                reminder_count = 0
            self.send_json(
                {
                    "model": ms,
                    "predicted_per_second": tps,
                    "overheated": oh,
                    "gpu_temp": gtemp,
                    "max_context": MAX_INPUT_TOKENS,
                    "reminder_count": reminder_count,
                }
            )
        elif self.path.startswith("/output/"):
            filename = os.path.basename(urlparse(self.path).path)
            fpath = os.path.abspath(os.path.join(COMFYUI_OUTPUT, filename))
            if fpath.startswith(os.path.abspath(COMFYUI_OUTPUT)) and os.path.exists(
                fpath
            ):
                self.send_response(200)
                self.send_header("Content-Type", "image/png")
                self.end_headers()
                with open(fpath, "rb") as f:
                    self.wfile.write(f.read())
                return
            self.send_error(404)
        elif self.path.startswith("/uploads/"):
            filename = os.path.basename(urlparse(self.path).path)
            fpath = os.path.abspath(os.path.join(UPLOADS_DIR, filename))
            if fpath.startswith(os.path.abspath(UPLOADS_DIR)) and os.path.exists(fpath):
                self.send_response(200)
                self.send_header("Content-Type", "application/octet-stream")
                self.send_header("Content-Disposition", "inline")
                self.end_headers()
                with open(fpath, "rb") as f:
                    self.wfile.write(f.read())
                return
            self.send_error(404)
        elif self.path.startswith("/api/status/"):
            task_id = os.path.basename(self.path)
            with _data_lock:
                status = tasks.get(
                    task_id, {"status": "unknown", "message": "Not found"}
                )
            self.send_json(status)
        elif self.path == "/api/sessions":
            user = get_current_user(self.headers)
            if not user:
                self.send_json([], status=401)
                return
            with _data_lock:
                sorted_items = sorted(
                    sessions_meta.items(),
                    key=lambda x: x[1].get("updated", 0),
                    reverse=True,
                )
                result = [
                    {
                        "session_id": sid,
                        "name": meta.get("name", "Chat"),
                        "created": meta.get("created", 0),
                        "updated": meta.get("updated", 0),
                        **context_token_report(sid, sessions.get(sid, [])),
                    }
                    for sid, meta in sorted_items
                    if meta.get("user_id", "") == user
                ]
            self.send_json(result)
        elif self.path.startswith("/api/sessions/") and self.path.endswith("/messages"):
            user = get_current_user(self.headers)
            if not user:
                self.send_json({"error": "Unauthorized"}, status=401)
                return
            sid = self.path.split("/")[3]
            with _data_lock:
                meta = sessions_meta.get(sid)
                if not meta or meta.get("user_id", "") != user:
                    self.send_error(404)
                    return
                msgs = sessions.get(sid)
            if msgs is not None:
                self.send_json(
                    {
                        "messages": msgs,
                        **context_token_report(sid, msgs),
                    }
                )
            else:
                self.send_error(404)
        elif self.path == "/":
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Cache-Control", "no-cache")
            self.end_headers()
            self.wfile.write(read_index_html().encode())
        else:
            DIST_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "dist")
            fpath = os.path.abspath(os.path.join(DIST_DIR, self.path.lstrip("/")))
            if fpath.startswith(os.path.abspath(DIST_DIR)) and os.path.isfile(fpath):
                ctype, _ = mimetypes.guess_type(fpath)
                self.send_response(200)
                self.send_header("Content-Type", ctype or "application/octet-stream")
                self.send_header("Cache-Control", "public, max-age=31536000, immutable")
                self.end_headers()
                with open(fpath, "rb") as f:
                    self.wfile.write(f.read())
            elif self.path.startswith("/api/") or "." in os.path.basename(self.path):
                if self.path == "/api/tasks":
                    user = get_current_user(self.headers)
                    if not user:
                        self.send_json({"error": "Unauthorized"}, status=401)
                        return
                    user_tasks = task_list(user)
                    self.send_json({"tasks": user_tasks})
                else:
                    self.send_error(404)
            else:
                self.send_response(200)
                self.send_header("Content-Type", "text/html; charset=utf-8")
                self.send_header("Cache-Control", "no-cache")
                self.end_headers()
                self.wfile.write(read_index_html().encode())

    def do_DELETE(self):
        if self.path.startswith("/api/sessions/"):
            user = get_current_user(self.headers)
            if not user:
                self.send_json({"error": "Unauthorized"}, status=401)
                return
            sid = self.path.split("/")[3]
            with _data_lock:
                meta = sessions_meta.get(sid)
                if not meta or meta.get("user_id", "") != user:
                    self.send_error(404)
                    return
                msgs = list(sessions.get(sid, []))
            for msg in msgs:
                if msg.get("role") == "assistant":
                    url = msg.get("_image_url", "") or ""
                    if url:
                        fname = os.path.basename(url)
                        fpath = os.path.join(IMG_PATH, fname)
                        if os.path.exists(fpath):
                            print(f"[delete] Removed output image: {fpath}")
                            os.remove(fpath)
                raw = msg.get("content", "")
                texts = []
                if isinstance(raw, str):
                    texts.append(raw)
                elif isinstance(raw, list):
                    for part in raw:
                        if isinstance(part, dict) and part.get("type") == "text":
                            texts.append(part.get("text", ""))
                for text in texts:
                    for part in text.split("[FILE:"):
                        idx = part.find("/uploads/")
                        if idx != -1:
                            url_part = part[idx:].split("]")[0]
                            fname = os.path.basename(url_part)
                            fpath = os.path.join(UPLOADS_DIR, fname)
                            if os.path.exists(fpath):
                                print(f"[delete] Removed uploaded file: {fpath}")
                                os.remove(fpath)

            with _data_lock:
                exists = sid in sessions
                if exists:
                    sessions.pop(sid, None)
                    sessions_meta.pop(sid, None)
            if exists:
                with _effective_contexts_lock:
                    _effective_contexts.pop(sid, None)
            if exists:
                save_sessions()
                self.send_json({"status": "deleted"})
            else:
                self.send_error(404)
        elif self.path.startswith("/api/tasks/"):
            user = get_current_user(self.headers)
            if not user:
                self.send_json({"error": "Unauthorized"}, status=401)
                return
            tid = self.path.split("/")[3]
            task_delete(tid, user)
            self.send_json({"status": "deleted"})
        else:
            self.send_error(404)

    def do_PUT(self):
        if self.path.startswith("/api/sessions/"):
            user = get_current_user(self.headers)
            if not user:
                self.send_json({"error": "Unauthorized"}, status=401)
                return
            sid = self.path.split("/")[3]
            length = int(self.headers.get("Content-Length", 0))
            body = json.loads(self.rfile.read(length))
            with _data_lock:
                meta = sessions_meta.get(sid)
                if meta and meta.get("user_id", "") == user:
                    meta["name"] = body.get("name", meta["name"])
                    meta["updated"] = time.time()
            if meta:
                save_sessions()
                self.send_json({"status": "updated"})
            else:
                self.send_error(404)
        elif self.path.startswith("/api/tasks/"):
            user = get_current_user(self.headers)
            if not user:
                self.send_json({"error": "Unauthorized"}, status=401)
                return
            tid = self.path.split("/")[3]
            length = int(self.headers.get("Content-Length", 0))
            body = json.loads(self.rfile.read(length))
            t = task_update(tid, user, **{k: v for k, v in body.items() if k in ("title","description","status","priority","due_date","reminder_at")})
            if t:
                self.send_json({"task": t})
            else:
                self.send_error(404)
        else:
            self.send_error(404)

    def do_POST(self):
        if self.path == "/api/login":
            length = int(self.headers.get("Content-Length", 0))
            body = json.loads(self.rfile.read(length))
            username = body.get("username", "")
            password = body.get("password", "")
            if get_user_password(username) == password:
                token = str(uuid.uuid4())
                with _tokens_lock:
                    _active_tokens[token] = username
                self.send_json(
                    {
                        "token": token,
                        "username": username,
                        "context_file": get_user_context_path(username),
                    }
                )
            else:
                self.send_json({"error": "Invalid credentials"}, status=401)
        elif self.path == "/api/logout":
            token = self.headers.get("X-Auth-Token", "")
            with _tokens_lock:
                _active_tokens.pop(token, None)
            self.send_json({"ok": True})
        elif self.path == "/api/user-context":
            user = get_current_user(self.headers)
            if not user:
                self.send_json({"error": "Unauthorized"}, status=401)
                return
            length = int(self.headers.get("Content-Length", 0))
            body = json.loads(self.rfile.read(length))
            action = body.get("action", "read")
            if action == "write":
                content = body.get("context", "")
                write_user_context(user, content)
                self.send_json({"status": "ok", "username": user})
            elif action == "overwrite":
                content = body.get("context", "")
                path = get_user_context_path(user)
                if path:
                    os.makedirs(os.path.dirname(path), exist_ok=True)
                    with open(path, "w") as f:
                        f.write(content)
                self.send_json({"status": "ok", "username": user})
            else:
                context = read_user_context(user)
                self.send_json(
                    {
                        "context": context,
                        "username": user,
                        "context_file": get_user_context_path(user),
                    }
                )
        elif self.path == "/api/chat":
            user = get_current_user(self.headers)
            if not user:
                self.send_json({"error": "Unauthorized"}, status=401)
                return
            length = int(self.headers.get("Content-Length", 0))
            body = json.loads(self.rfile.read(length))
            task_id = str(uuid.uuid4())
            sid = body.get("session_id", "default")
            with _data_lock:
                meta = sessions_meta.get(sid)
                if not meta or meta.get("user_id", "") != user:
                    self.send_json({"error": "Session not found"}, status=404)
                    return

            entry = {
                "task_id": task_id,
                "session_id": sid,
                "message": body.get("message", ""),
                "image": body.get("image"),
                "audio": body.get("audio"),
                "user": user,
                "client_timestamp": body.get("client_timestamp"),
            }
            with _queue_lock:
                if len(_task_queue) >= MAX_QUEUE_SIZE:
                    self.send_json({"error": "Server busy"}, status=503)
                    return
                _task_queue.append(entry)
                _queue_cond.notify()
            with _data_lock:
                tasks[task_id] = {
                    "status": "queued",
                    "message": "Waiting in line...",
                    "session_id": sid,
                }
            self.send_json({"task_id": task_id})
        elif self.path == "/api/extract-file":
            length = int(self.headers.get("Content-Length", 0))
            body = json.loads(self.rfile.read(length))
            name = body.get("name", "")
            data_b64 = body.get("data", "")
            ext = os.path.splitext(name)[1].lower()
            safe_name = str(uuid.uuid4()) + ext
            filepath = os.path.join(UPLOADS_DIR, safe_name)
            raw = base64.b64decode(data_b64)
            with open(filepath, "wb") as f:
                f.write(raw)
            file_url = f"/uploads/{safe_name}"
            self.send_json({"url": file_url, "name": name})
        elif self.path == "/api/tts":
            user = get_current_user(self.headers)
            if not user:
                self.send_json({"error": "Unauthorized"}, status=401)
                return
            length = int(self.headers.get("Content-Length", 0))
            body = json.loads(self.rfile.read(length))
            raw_text = body.get("text", "")
            if not raw_text:
                self.send_json({"error": "No text provided"}, status=400)
                return
            try:
                import re
                text = raw_text
                voice = body.get("voice", "")

                # Detect language tag from LLM prefix: [bn], [hi], [te], [kn], [en]
                m = re.match(r"^\s*\[(bn|hi|te|kn|en)\]\s*", text)
                if m:
                    tag = m.group(1)
                    text = text[m.end():]
                elif not voice:
                    bn = len(re.findall(r"[\u0980-\u09FF]", text))
                    hi = len(re.findall(r"[\u0900-\u097F]", text))
                    te = len(re.findall(r"[\u0C00-\u0C7F]", text))
                    kn = len(re.findall(r"[\u0C80-\u0CFF]", text))
                    scores = {"bn": bn, "hi": hi, "te": te, "kn": kn}
                    tag = max(scores, key=scores.get)
                    if scores[tag] == 0:
                        tag = "en"
                else:
                    tag = "en"

                # Determine TTS backend
                PIPER_VOICES = {
                    "bn": "/home/palash/.piper_voices/bn_BD-google-medium.onnx",
                    "hi": "/home/palash/.piper_voices/hi_IN-priyamvada-medium.onnx",
                    "te": "/home/palash/.piper_voices/te_IN-padmavathi-medium.onnx",
                    "en": "/home/palash/.piper_voices/en_US-amy-medium.onnx",
                }
                EDGE_VOICES = {
                    "bn": "bn-IN-TanishaaNeural",
                    "hi": "hi-IN-SwaraNeural",
                    "te": "te-IN-ShrutiNeural",
                    "kn": "kn-IN-GaganNeural",
                    "en": "en-US-AriaNeural",
                }

                if tag in PIPER_VOICES:
                    import piper, io, struct, wave
                    onnx_path = PIPER_VOICES[tag]
                    cfg_path = onnx_path + ".json"
                    if not hasattr(self, "_piper_voices"):
                        self._piper_voices = {}
                    if tag not in self._piper_voices:
                        print(f"[tts] Loading Piper voice '{tag}' ...")
                        self._piper_voices[tag] = piper.PiperVoice.load(
                            onnx_path, config_path=cfg_path
                        )
                    pv = self._piper_voices[tag]
                    print(f"[tts] Piper {tag}: synthesizing {len(text)} chars")
                    wav_io = io.BytesIO()
                    with wave.open(wav_io, "wb") as wf:
                        wf.setnchannels(1)
                        wf.setsampwidth(2)
                        wf.setframerate(22050)
                        for chunk in pv.synthesize(text):
                            int16 = (chunk.audio_float_array * 32767).clip(-32768, 32767).astype("<i2")
                            wf.writeframes(int16.tobytes())
                    audio_b64 = base64.b64encode(wav_io.getvalue()).decode()
                    self.send_json({"audio": audio_b64, "type": "audio/wav"})
                else:
                    import asyncio, edge_tts
                    edge_voice = voice or EDGE_VOICES.get(tag, "en-US-AriaNeural")
                    print(f"[tts] edge-tts {tag} ({edge_voice}): {len(text)} chars")
                    communicate = edge_tts.Communicate(text, edge_voice)
                    mp3_data = bytearray()
                    async def _gen():
                        async for chunk in communicate.stream():
                            if chunk["type"] == "audio":
                                mp3_data.extend(chunk["data"])
                    asyncio.run(_gen())
                    audio_b64 = base64.b64encode(bytes(mp3_data)).decode()
                    self.send_json({"audio": audio_b64, "type": "audio/mpeg"})
            except Exception as e:
                print(f"[tts] Error: {e}")
                traceback.print_exc()
                self.send_json({"error": str(e)}, status=500)
        elif self.path == "/api/sessions":
            user = get_current_user(self.headers)
            if not user:
                self.send_json({"error": "Unauthorized"}, status=401)
                return
            sid = str(uuid.uuid4())
            now = time.time()
            with _data_lock:
                sessions[sid] = []
                sessions_meta[sid] = {
                    "name": "New Chat",
                    "created": now,
                    "updated": now,
                    "user_id": user,
                }
            save_sessions()
            self.send_json({"session_id": sid})
        elif self.path == "/api/location":
            length = int(self.headers.get("Content-Length", 0))
            body = json.loads(self.rfile.read(length))
            global _client_location
            task_id = body.get("task_id")
            if body.get("denied"):
                _client_location = ""
                ev = _location_events.get(task_id) if task_id else None
                if ev:
                    ev.set()
                self.send_json({"ok": True})
                return
            lat = body.get("latitude")
            lng = body.get("longitude")
            if lat is not None and lng is not None:
                try:
                    geo = requests.get(
                        "https://nominatim.openstreetmap.org/reverse",
                        params={"format": "json", "lat": lat, "lon": lng},
                        headers={"User-Agent": "LocalAI/1.0"},
                        timeout=5,
                    ).json()
                    display = geo.get("display_name", "")
                    _client_location = display
                except Exception:
                    _client_location = f"{lat:.4f}, {lng:.4f}"
            ev = _location_events.get(task_id) if task_id else None
            if ev:
                ev.set()
            self.send_json({"ok": True})
        elif self.path == "/api/tasks":
            user = get_current_user(self.headers)
            if not user:
                self.send_json({"error": "Unauthorized"}, status=401)
                return
            length = int(self.headers.get("Content-Length", 0))
            body = json.loads(self.rfile.read(length))
            t = task_create(user, body.get("title", "Untitled"), body.get("description", ""), body.get("priority", "medium"), body.get("due_date"), body.get("session_id"), body.get("reminder_at"))
            self.send_json({"task": t})
        else:
            self.send_error(404)

    def send_json(self, data, status=200):
        self.send_response(status)
        self.send_header("Content-Type", "application/json")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.end_headers()
        self.wfile.write(json.dumps(data).encode())

    def log_message(self, format, *args):
        pass


if __name__ == "__main__":
    os.makedirs(UPLOADS_DIR, exist_ok=True)
    load_sessions()
    try:
        r = requests.get(f"{LLAMA_BASE}/health", timeout=3)
        if r.status_code != 200:
            raise Exception("health check failed")
        print("[startup] llama-server is running")
    except Exception:
        print("[startup] llama-server not reachable — starting...")
        restart_servers()
    try:
        r = requests.get(SEARXNG_URL, timeout=3)
        if r.status_code in (200, 301, 302):
            print("[startup] SearXNG is running")
        else:
            raise Exception(f"status {r.status_code}")
    except Exception as e:
        print(f"[startup] ERROR: SearXNG is not reachable at {SEARXNG_URL} ({e}). Web search will not work. Exiting.")
        sys.exit(1)
    threading.Thread(target=_event_loop, daemon=True).start()
    threading.Thread(target=_queue_worker, daemon=True).start()
    threading.Thread(target=_idle_unload_loop, daemon=True).start()
    threading.Thread(target=_thermal_monitor, daemon=True).start()
    threading.Thread(target=_reminder_loop, daemon=True).start()
    print(f"Chat UI running on http://localhost:{PORT}")
    s = http.server.HTTPServer((HOST, PORT), Handler)
    try:
        s.serve_forever()
    except KeyboardInterrupt:
        s.shutdown()
